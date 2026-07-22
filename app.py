from pathlib import Path
content = '''# app.py
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
import streamlit as st

APP_TITLE = "Seating Plan Generator"
DEFAULT_LEFT_5 = [8, 6, 10, 12, 14]
DEFAULT_CENTER_5 = [2, 1, 3, 4, 5]
DEFAULT_RIGHT_5 = [9, 7, 11, 13, 15]
DEFAULT_CENTER_6 = [4, 1, 2, 3, 5, 6]

CENTER_5 = {
    "top": [2, 1, 3],
    "bottom": [4, 5],
}
CENTER_6 = {
    "top": [4, 1, 2, 3],
    "bottom": [5, 6],
}
SIDE_5_LEFT = {"top": [8, 6, 10], "bottom": [12, 14]}
SIDE_5_RIGHT = {"top": [9, 7, 11], "bottom": [13, 15]}


def set_cell_text(cell, text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    return p


def set_cell_shading(cell, fill: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_cell_margins(cell, top=60, start=60, bottom=60, end=60):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color="000000", size=8):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        element = tcBorders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tcBorders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def merge_span(table, row_idx, start_col, end_col, text, bold=False, size=11):
    cell = table.cell(row_idx, start_col)
    if end_col > start_col:
        cell = cell.merge(table.cell(row_idx, end_col))
    set_cell_text(cell, text, bold=bold, size=size)
    set_cell_shading(cell, "E9E2C7")
    set_cell_border(cell, size=8)
    set_cell_margins(cell, 80, 80, 80, 80)
    return cell


def write_seat_cell(cell, seat_no: int, code: str):
    set_cell_text(cell, seat_no, bold=True, size=11)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(code)
    run.bold = True
    run.font.size = Pt(8)
    run.font.name = "Arial"
    set_cell_shading(cell, "F5EFD6")
    set_cell_border(cell, size=5)
    set_cell_margins(cell, 80, 80, 80, 80)


def landscape(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)


def add_title(doc: Document, title: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "Arial"


def add_three_table_layout(doc: Document, group_map: Dict[str, List[Tuple[int, str]]]):
    trio = doc.add_table(rows=1, cols=3)
    trio.alignment = WD_TABLE_ALIGNMENT.CENTER
    trio.autofit = False
    trio.allow_autofit = False
    trio.columns[0].width = Inches(3.0)
    trio.columns[1].width = Inches(3.3)
    trio.columns[2].width = Inches(3.0)

    templates = {
        5: {"n_cols": 3, "top": [2, 1, 3], "bottom": [4, 5], "label_cols": (1, 1)},
        6: {"n_cols": 4, "top": [4, 1, 2, 3], "bottom": [5, 6], "label_cols": (1, 2)},
        7: {"n_cols": 4, "top": [4, 1, 2, 3], "bottom": [5, 6, 7], "label_cols": (1, 2)},
        8: {"n_cols": 4, "top": [4, 1, 2, 3], "bottom": [5, 6, 7, 8], "label_cols": (1, 2)},
    }

    for idx, grp in enumerate(["Left", "Center", "Right"]):
        seats = group_map.get(grp, [])
        n = min(len(seats), 8)
        tmpl = templates.get(n, templates[6])
        cell = trio.cell(0, idx)
        cell.text = ""
        cell.vertical_alignment = 0
        spacer = cell.paragraphs[0]
        spacer.text = ""
        label = cell.add_paragraph()
        label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = label.add_run(f"{grp} Table")
        rr.bold = True
        rr.font.size = Pt(10)
        rr.font.name = "Arial"
        rr.font.color.rgb = None

        table = cell.add_table(rows=2, cols=tmpl["n_cols"])
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.allow_autofit = False
        widths = [0.82] * tmpl["n_cols"]
        widths[tmpl["n_cols"] // 2] = 1.05 if grp == "Center" else 0.9
        for c in range(tmpl["n_cols"]):
            for r in range(2):
                table.cell(r, c).width = Inches(widths[c])
                table.cell(r, c).text = ""
                set_cell_border(table.cell(r, c), size=5)

        lookup = {seat_no: code for seat_no, code in seats}
        for seat_no, col in zip(tmpl["top"], range(len(tmpl["top"]))):
            write_seat_cell(table.cell(0, col), seat_no, lookup.get(seat_no, ""))
        for seat_no, col in zip(tmpl["bottom"], [0, tmpl["n_cols"] - 1]):
            write_seat_cell(table.cell(1, col), seat_no, lookup.get(seat_no, ""))

        a, b = tmpl["label_cols"]
        label_cell = table.cell(1, a)
        if a != b:
            label_cell = label_cell.merge(table.cell(1, b))
        set_cell_text(label_cell, "MAIN" if grp == "Center" else grp, bold=True, size=12 if grp == "Center" else 11)
        set_cell_shading(label_cell, "E9E2C7")
        set_cell_border(label_cell, size=8)
        set_cell_margins(label_cell, 80, 80, 80, 80)


def build_doc(df: pd.DataFrame) -> Document:
    doc = Document()
    landscape(doc)
    add_title(doc, "Inaugural Seating Plan")
    group_map = {
        "Left": [(int(r.seat_no), str(r.code)) for r in df[df["group"] == "Left"].itertuples(index=False)],
        "Center": [(int(r.seat_no), str(r.code)) for r in df[df["group"] == "Center"].itertuples(index=False)],
        "Right": [(int(r.seat_no), str(r.code)) for r in df[df["group"] == "Right"].itertuples(index=False)],
    }
    add_three_table_layout(doc, group_map)
    return doc


def main():
    st.set_page_config(page_title=APP_TITLE, layout="wide")
    st.title(APP_TITLE)
    st.caption("Upload CSV with columns: seat_no, code, group")
    uploaded = st.file_uploader("CSV", type=["csv"])
    if uploaded:
        df = pd.read_csv(uploaded)
        if not {"seat_no", "code", "group"}.issubset(df.columns):
            st.error("CSV must contain seat_no, code, group")
            return
        doc = build_doc(df)
        out = Path("output")
        out.mkdir(exist_ok=True)
        path = out / "app_generated.docx"
        doc.save(path)
        st.success(f"Saved {path}")
        with open(path, "rb") as f:
            st.download_button("Download DOCX", f, file_name="app_generated.docx")
    else:
        st.info("Upload CSV to generate the Word file.")


if __name__ == "__main__":
    main()
'''
Path('output/app.py').write_text(content)
print('full app.py written')
